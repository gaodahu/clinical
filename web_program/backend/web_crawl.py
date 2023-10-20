import time
from selenium import webdriver
from time import sleep
from selenium.webdriver import ChromeOptions
from selenium.webdriver.common.keys import Keys
import pandas as pd
from selenium.common.exceptions import NoSuchElementException
from docx import Document
from docx.shared import Pt
import os
import transformers
import bitsandbytes
from transformers import BartTokenizer
from transformers import AutoTokenizer
from transformers import FlaxBartForConditionalGeneration
from transformers import BartForConditionalGeneration
import accelerate
import torch
import typing_extensions
from llama_index import VectorStoreIndex, SimpleDirectoryReader, ServiceContext
from llama_index.llms import HuggingFaceLLM
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from llama_index import LangchainEmbedding, ServiceContext
import llama_index
from llama_index.prompts.prompts import SimpleInputPrompt
import logging
import sys
import glob

os.environ["HUGGINGFACE_TOKEN"] = "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"
def clear_directory(directory_path):
    # Get a list of all the file paths in the directory
    files = glob.glob(directory_path + '/*')

    # Remove each file
    for file in files:
        try:
            os.remove(file)
        except Exception as e:
            print(f"Error deleting file {file}: {e}")


def is_element_present(by, value):
    try:
        element = driver.find_element(by=by, value=value)
    except NoSuchElementException as e:
        return False
    return True
def standardize_quotes(s):
    s = s.replace('“', '"').replace('”', '"')
    s = s.replace('‘', "'").replace('’', "'")
    s = s.replace('—', '-')
    return s

def add_text_with_formatting(document, text):
    start = 0
    while True:
        format_start = text.find('**', start)
        if format_start == -1:
            break
        p = document.add_paragraph(text[start:format_start])
        set_font_for_paragraph(p, 'Montserrat', 12)

        format_end = text.find('**', format_start + 2)
        if format_end == -1:
            p = document.add_paragraph(text[format_start:])
            set_font_for_paragraph(p, 'Montserrat', 12)
            break

        p = document.add_paragraph()
        run = p.add_run(text[format_start + 2:format_end])
        run.bold = True
        run.underline = True
        set_font_for_run(run, 'Montserrat', 12)
        start = format_end + 2

    p = document.add_paragraph(text[start:])
    set_font_for_paragraph(p, 'Montserrat', 12)

def set_font_for_paragraph(paragraph, font_name, size):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)


def set_font_for_run(run, font_name, size):
    run.font.name = font_name
    run.font.size = Pt(size)



input_text = r"SUBJECT: CPTII Deficiency MESSAGE: My sister in Texas was just diagnosed with CPTII deficiency today.  We do not yet know which variant.  I was told a year ago following her near death from rhabdo (CK's 150,000+), ATN, RF, respiratory compromise, to be tested for a FA metabolic error.  Until today, I really had no idea which test to request.  Her muscle bx was performed a few months ago as she was still ill last year.  Another recurrence of rhado this year prompted the bx (she is retired military).  I don't know which branch of medicine to contact for genetic testing. From  the NIH website, muscle bx is the most definitive.  I also found a phone number for the Genetics and Rare Diseases Information Center.  Thank you for some direction."

def generate_response(input_text):
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    data = pd.read_excel('/mnt/d/data.xlsx')

    model_path = '/mnt/d/bart_results/bart_results/checkpoint-2500'
    model_name = "/mnt/d/bart/bart"
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model = BartForConditionalGeneration.from_pretrained(model_path)
    tokenizer = BartTokenizer.from_pretrained(model_name)
    model = model.to(device)

    inputs = tokenizer(input_text, return_tensors='pt')
    inputs = {name: tensor.to(model.device) for name, tensor in inputs.items()}
    outputs = model.generate(**inputs)
    prediction = tokenizer.decode(outputs[0], skip_special_tokens=True)

    logging.basicConfig(stream=sys.stdout, level=logging.INFO)
    logging.getLogger().addHandler(logging.StreamHandler(stream=sys.stdout))

    options = webdriver.ChromeOptions()
    options.add_argument("--headless")
    options.add_argument("--enable-gpu-rasterization")
    options.add_experimental_option('excludeSwitches', ['enable-automation'])
    driver=webdriver.Chrome('/mnt/d/chromedriver-linux64/chromedriver', options=options)

    # driver = webdriver.Remote(command_executor="http://localhost:9515", options=options)
    driver.get("https://www.google.com")

    directory_path = "/mnt/d/project1/supportingdocument/"
    clear_directory(directory_path)

    url='https://www.google.com/'
    driver.get(url)
    torch.cuda.empty_cache()
    search_textbox=driver.find_element('xpath','//textarea[@class="gLFyf"]')
    search_textbox.clear()
    search_textbox.send_keys(prediction)
    search_textbox.send_keys(Keys.ENTER)
    sleep(4)
    Web_links=driver.find_elements('xpath','//div[@class="yuRUbf"]//h3//..//../a[@href]')
    webpage_links = []
    for element in Web_links:
        webpage_links.append(element.get_attribute('href'))
    index=0
    for link in webpage_links[:3]:
        sleep(2)
        driver.get(link)
        crawled_texts = []
        paragraphes=driver.find_elements('xpath','//p')
        for element in paragraphes:
                crawled_texts.append(element.text)
        filtered_lst_crawled_texts = [item for item in crawled_texts if item]
        filtered_lst_crawled_texts = [sentence for sentence in filtered_lst_crawled_texts if len(sentence.split()) >= 3]
        combined_string_crawled_texts = '\n'.join(filtered_lst_crawled_texts)
        combined_string_crawled_texts = combined_string_crawled_texts.replace('\n\n', '\n')
        combined_string_crawled_texts = standardize_quotes(combined_string_crawled_texts)
        filename = "/mnt/d/project1/supportingdocument/" + str(index) + ".doc"
        doc = Document()
        add_text_with_formatting(doc, combined_string_crawled_texts)
        doc.save(filename)
        index+=1

    system_prompt = "You are a clinical consultant. Your goal is to answer questions as accurately as possible based on the instructions and context provided."

    query_wrapper_prompt = SimpleInputPrompt("<|USER|>{query_str}<|ASSISTANT|>")



    llm = HuggingFaceLLM(
        context_window=1024,
        max_new_tokens=128,
        generate_kwargs={"temperature": 0.0, "do_sample": False},
        system_prompt=system_prompt,
        query_wrapper_prompt=query_wrapper_prompt,
        tokenizer_name="meta-llama/Llama-2-7b-chat-hf",
        model_name="meta-llama/Llama-2-7b-chat-hf",
        device_map="auto",
        model_kwargs={"torch_dtype": torch.float16 }
    )

    prediction='what is diabetes?'

    from sentence_transformers import SentenceTransformer, util

    # Setting the device explicitly
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    # Load the model and push to the device
    model_name = "sentence-transformers/all-mpnet-base-v2"
    sentence_transformer_model = SentenceTransformer(model_name, device=device)

    embed_model = LangchainEmbedding(
        HuggingFaceEmbeddings(model_name=model_name)
    )

    # embed_model = LangchainEmbedding(
    #   HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    # )

    service_context = ServiceContext.from_defaults(
        chunk_size=256,
        llm=llm,
        embed_model=embed_model
    )

    documents = SimpleDirectoryReader("/mnt/d/project1/supportingdocument/").load_data()
    # from llama_index import Document
    # document = Document(
    #     text='text',
    #     metadata={
    #         'filename': '<doc_file_name>',
    #     }
    # )
    # document.metadata = {'filename': '<doc_file_name>'}
    # from llama_index import SimpleDirectoryReader
    # filename_fn = lambda filename: {'file_name': filename}
    # documents = SimpleDirectoryReader("/mnt/d/project1/supportingdocument/", file_metadata=filename_fn).load_data()
    # document.excluded_llm_metadata_keys = ['file_name']
    # from llama_index.schema import MetadataMode
    # print(document.get_content(metadata_mode=MetadataMode.LLM))
    # document.excluded_embed_metadata_keys = ['file_name']
    # from llama_index.schema import MetadataMode
    # print(document.get_content(metadata_mode=MetadataMode.EMBED))


    index = VectorStoreIndex.from_documents(documents, service_context=service_context)

    query_engine = index.as_query_engine()
    response = query_engine.query(prediction)
    str(response.metadata)
    # find = re.findall(r"'page_label': '[^']*', 'file_name': '[^']*'", document_info)
    print(response)
    return response
response=generate_response(input_text)
str(response)
print ((str(response)))
json_string='{"response":'+'"'+str(response)+'"}'
json_string = '{"response": "Diabetes is a chronic medical condition characterized by high blood sugar levels, which can lead to serious health complications if left untreated. There are two main types of diabetes: type 1 and type 2. Type 1 diabetes is an autoimmune disease in which the body\'s immune system attacks and destroys the cells in the pancreas that produce insulin, a hormone that regulates blood sugar levels. Type 2 diabetes is the most common form of diabetes, and it occurs when the body becomes resistant to insulin."}'

import json
json_object = json.loads(json_string)