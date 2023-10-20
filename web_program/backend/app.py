from flask import Flask, request, jsonify
import time
from selenium import webdriver
from time import sleep
from selenium.webdriver import ChromeOptions
from selenium.webdriver.common.keys import Keys
import pandas as pd
from selenium.common.exceptions import NoSuchElementException
from docx import Document
from docx.shared import Pt
import os
import transformers
import bitsandbytes
from transformers import BartTokenizer
from transformers import AutoTokenizer
from transformers import FlaxBartForConditionalGeneration
from transformers import BartForConditionalGeneration
import accelerate
import torch
import typing_extensions
from llama_index import VectorStoreIndex, SimpleDirectoryReader, ServiceContext
from llama_index.llms import HuggingFaceLLM
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from llama_index import LangchainEmbedding, ServiceContext
import llama_index
from llama_index.prompts.prompts import SimpleInputPrompt
import logging
import sys
import glob
from flask_socketio import SocketIO, emit
from transformers import XLNetForSequenceClassification, XLNetTokenizer
import random
os.environ["HUGGINGFACE_TOKEN"] = "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"
def clear_directory(directory_path):
    # Get a list of all the file paths in the directory
    files = glob.glob(directory_path + '/*')

    # Remove each file
    for file in files:
        try:
            os.remove(file)
        except Exception as e:
            print(f"Error deleting file {file}: {e}")


def is_element_present(by, value):
    try:
        element = driver.find_element(by=by, value=value)
    except NoSuchElementException as e:
        return False
    return True
def standardize_quotes(s):
    s = s.replace('“', '"').replace('”', '"')
    s = s.replace('‘', "'").replace('’', "'")
    s = s.replace('—', '-')
    return s

def add_text_with_formatting(document, text):
    start = 0
    while True:
        format_start = text.find('**', start)
        if format_start == -1:
            break
        p = document.add_paragraph(text[start:format_start])
        set_font_for_paragraph(p, 'Montserrat', 12)

        format_end = text.find('**', format_start + 2)
        if format_end == -1:
            p = document.add_paragraph(text[format_start:])
            set_font_for_paragraph(p, 'Montserrat', 12)
            break

        p = document.add_paragraph()
        run = p.add_run(text[format_start + 2:format_end])
        run.bold = True
        run.underline = True
        set_font_for_run(run, 'Montserrat', 12)
        start = format_end + 2

    p = document.add_paragraph(text[start:])
    set_font_for_paragraph(p, 'Montserrat', 12)

def set_font_for_paragraph(paragraph, font_name, size):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)


def set_font_for_run(run, font_name, size):
    run.font.name = font_name
    run.font.size = Pt(size)

logging.basicConfig(stream=sys.stdout, level=logging.INFO)
logging.getLogger().addHandler(logging.StreamHandler(stream=sys.stdout))

options = webdriver.ChromeOptions()
options.add_argument("--headless")
options.add_argument("--enable-gpu-rasterization")
options.add_experimental_option('excludeSwitches', ['enable-automation'])
driver=webdriver.Chrome('/mnt/d/chromedriver-linux64/chromedriver', options=options)
system_prompt = "You are a clinical consultant. Your goal is to answer questions as accurately as possible based on the instructions and context provided."

query_wrapper_prompt = SimpleInputPrompt("<|USER|>{query_str}<|ASSISTANT|>")
# Setting the device explicitly
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model_path = '/mnt/d/bart_results/bart_results/checkpoint-2500'
model_name = "/mnt/d/bart/bart"
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model = BartForConditionalGeneration.from_pretrained(model_path)
tokenizer = BartTokenizer.from_pretrained(model_name)
model = model.to(device)


llm = HuggingFaceLLM(
    context_window=1024,
    max_new_tokens=128,
    generate_kwargs={"temperature": 0.0, "do_sample": False},
    system_prompt=system_prompt,
    query_wrapper_prompt=query_wrapper_prompt,
    tokenizer_name="meta-llama/Llama-2-7b-chat-hf",
    model_name="meta-llama/Llama-2-7b-chat-hf",
    device_map="auto",
    model_kwargs={"torch_dtype": torch.float16 }
)

from sentence_transformers import SentenceTransformer, util


# Load the model and push to the device
model_name = "sentence-transformers/all-mpnet-base-v2"
sentence_transformer_model = SentenceTransformer(model_name, device=device)

embed_model = LangchainEmbedding(
    HuggingFaceEmbeddings(model_name=model_name)
)

service_context = ServiceContext.from_defaults(
    chunk_size=256,
    llm=llm,
    embed_model=embed_model
)


# input_text = r"SUBJECT: CPTII Deficiency MESSAGE: My sister in Texas was just diagnosed with CPTII deficiency today.  We do not yet know which variant.  I was told a year ago following her near death from rhabdo (CK's 150,000+), ATN, RF, respiratory compromise, to be tested for a FA metabolic error.  Until today, I really had no idea which test to request.  Her muscle bx was performed a few months ago as she was still ill last year.  Another recurrence of rhado this year prompted the bx (she is retired military).  I don't know which branch of medicine to contact for genetic testing. From  the NIH website, muscle bx is the most definitive.  I also found a phone number for the Genetics and Rare Diseases Information Center.  Thank you for some direction."
# answer=generate_response(input_text)
def entity_recognition():
    import xml.etree.ElementTree as ET

    def parse_xml(xml_file):
        # Parse XML with ElementTree
        tree = ET.parse(xml_file)
        root = tree.getroot()

        text_list = []

        # Iterate through each element in the XML
        for elem in root.iter('String'):
            text_list.append(elem.text)

        return text_list

    qual2023 = '/content/drive/MyDrive/xml/qual2023.xml'
    pa2023 = '/content/drive/MyDrive/xml/pa2023.xml'
    desc2023 = '/content/drive/MyDrive/xml/desc2023.xml'
    supp2023 = '/content/drive/MyDrive/xml/supp2023.xml'
    qual2023 = parse_xml(qual2023)
    pa2023 = parse_xml(pa2023)
    desc2023 = parse_xml(desc2023)
    supp2023 = parse_xml(supp2023)
    combined = qual2023 + pa2023 + desc2023 + qual2023
    unique_words = list(set(combined))
    drug_path = '/content/drive/MyDrive/drug/Products.txt'
    import pandas as pd
    data = pd.read_csv(drug_path, delimiter='\t', error_bad_lines=False)
    import pandas as pd

    # Load the data from the text file
    drug_path = '/content/drive/MyDrive/drug/Products.txt'

    data = pd.read_csv(drug_path, delimiter='\t', error_bad_lines=False)
    drug_names = data['DrugName'].tolist()
    unique_drugs = list(set(drug_names))
    combined = unique_drugs + unique_words

    # Convert to a set to remove duplicates, then convert back to a list
    vocabulary = list(set(combined))
    return vocabulary

def information_extraction(text):
    import torch
    model_path = "/mnt/d/xlnet"
    model = XLNetForSequenceClassification.from_pretrained(model_path)
    tokenizer = XLNetTokenizer.from_pretrained(model_path)
    inputs = tokenizer(text, return_tensors="pt", max_length=1024, truncation=True)
    model.eval()
    with torch.no_grad():
        outputs = model(**inputs)
    logits = outputs.logits
    predictions = torch.argmax(logits, dim=-1)
    Q_type= str(predictions.item())
    return Q_type


def replace_first_word(input_string, new_word):
    words = input_string.split(' ')
    words[0] = new_word
    return ' '.join(words)
def generate_response(input_text,information_extract):
    print("Received POST request!")
    socketio.emit('response', {'data': 'Started Processing'}, namespace='/status')
    # device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    socketio.emit('response', {'data': 'BART Processing'}, namespace='/status')
    inputs = tokenizer(input_text, return_tensors='pt')
    inputs = {name: tensor.to(model.device) for name, tensor in inputs.items()}
    outputs = model.generate(**inputs)
    prediction = tokenizer.decode(outputs[0], skip_special_tokens=True)
    if information_extract is 'on':
        Q_type=information_extraction(input_text)
        prediction=replace_first_word(prediction,Q_type)
    else:
        pass
    socketio.emit('response', {'data': 'Crawling'}, namespace='/status')
    # driver = webdriver.Remote(command_executor="http://localhost:9515", options=options)
    driver.get("https://www.google.com")

    directory_path = "/mnt/d/project1/supportingdocument/"
    clear_directory(directory_path)

    url='https://www.google.com/'
    driver.get(url)
    torch.cuda.empty_cache()
    search_textbox=driver.find_element('xpath','//textarea[@class="gLFyf"]')
    search_textbox.clear()
    search_textbox.send_keys(prediction)
    search_textbox.send_keys(Keys.ENTER)
    sleep(4)
    Web_links=driver.find_elements('xpath','//div[@class="yuRUbf"]//h3//..//../a[@href]')
    webpage_links = []
    for element in Web_links:
        webpage_links.append(element.get_attribute('href'))
    index=0

    for link in webpage_links[:3]:
        sleep(2)
        driver.get(link)
        crawled_texts = []
        paragraphes=driver.find_elements('xpath','//p')
        for element in paragraphes:
                crawled_texts.append(element.text)
        filtered_lst_crawled_texts = [item for item in crawled_texts if item]
        filtered_lst_crawled_texts = [sentence for sentence in filtered_lst_crawled_texts if len(sentence.split()) >= 3]
        combined_string_crawled_texts = '\n'.join(filtered_lst_crawled_texts)
        combined_string_crawled_texts = combined_string_crawled_texts.replace('\n\n', '\n')
        combined_string_crawled_texts = standardize_quotes(combined_string_crawled_texts)
        filename = "/mnt/d/project1/supportingdocument/" + str(index) + ".doc"
        doc = Document()
        add_text_with_formatting(doc, combined_string_crawled_texts)
        doc.save(filename)
        index+=1
    socketio.emit('response', {'data': 'Indexing'}, namespace='/status')
    documents = SimpleDirectoryReader("/mnt/d/project1/supportingdocument/").load_data()

    index = VectorStoreIndex.from_documents(documents, service_context=service_context)
    socketio.emit('response', {'data': 'GeneratingAnswer'}, namespace='/status')
    query_engine = index.as_query_engine()
    #https://gpt-index.readthedocs.io/en/latest/core_modules/data_modules/documents_and_nodes/usage_documents.html
    response = query_engine.query(prediction)
    print(response)
    response=str(response)
    return response, prediction
def read_example_input():
    data = pd.read_excel('/mnt/d/data.xlsx')
    input_list= data['CHQ']
    selected_elements = input_list.sample(3).reset_index(drop=True)
    example1 = selected_elements[0]
    example2 = selected_elements[1]
    example3 = selected_elements[2]
    return example1,example2,example3


from flask_cors import CORS
app = Flask(__name__)

socketio = SocketIO(app, cors_allowed_origins="*")

@socketio.on('connect', namespace='/status')
def handle_connect():
    print('Client Connected')
    emit('response', {'data': 'connected'})
@socketio.on('disconnect', namespace='/status')
def handle_disconnect():
    print('Client Disconnected')

CORS(app)

@app.route('/api/ask', methods=['POST'])
def process_question():
    input_text = request.json.get('question')
    information_extract = request.json.get('enable_entity_recognition', 'off')
    response, prediction = generate_response(input_text,information_extract)
    return jsonify({"response": response, "prediction": prediction})

@app.route('/api/example', methods=['GET'])
def generate_example():
    example_1, example_2,example_3 = read_example_input()
    print('example created')
    return jsonify({"example_1": example_1, "example_2": example_2, "example_3": example_3})


if __name__ == '__main__':
    socketio.run(app, debug=True, allow_unsafe_werkzeug=True)